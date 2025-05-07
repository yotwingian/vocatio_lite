import streamlit as st
import os
import tempfile
import openai
from pathlib import Path
import docx
from docx import Document
import base64
import subprocess
import shutil
import torch
from diarization import Diarization

# Try to load environment variables from .env file (optional dependency)
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    # python-dotenv is optional; continue if not available
    pass

# Set page configuration
st.set_page_config(page_title="Vocatio Lite", page_icon="🎤", layout="wide")

# Initialize session state variables
if "transcript" not in st.session_state:
    st.session_state.transcript = None
if "transcription_completed" not in st.session_state:
    st.session_state.transcription_completed = False
if "file_name" not in st.session_state:
    st.session_state.file_name = None
if "diarization_result" not in st.session_state:
    st.session_state.diarization_result = None
if "aligned_transcript" not in st.session_state:
    st.session_state.aligned_transcript = None

def get_openai_key():
    """Get the OpenAI API key from environment variables"""
    return os.environ.get("OPENAI_API_KEY")

def get_hf_token():
    """Get the HuggingFace access token from environment variables"""
    return os.environ.get("HF_AUTH_TOKEN")

def compress_audio(input_path):
    """
    Compress audio file to reduce size while maintaining quality for transcription
    
    Args:
        input_path: Path to the input audio file
        
    Returns:
        Path to the compressed audio file
    """
    # Create a temporary file with .ogg extension
    temp_dir = tempfile.gettempdir()
    output_path = os.path.join(temp_dir, f"compressed_{Path(input_path).stem}.ogg")
    
    try:
        # Check if ffmpeg is available
        try:
            subprocess.run(["ffmpeg", "-version"], check=True, capture_output=True)
        except (subprocess.SubprocessError, FileNotFoundError):
            st.warning("ffmpeg not found. Processing without compression.")
            return input_path
            
        # Compress audio using ffmpeg (convert to mono Opus with 12kbps bitrate)
        subprocess.run([
            "ffmpeg", 
            "-i", input_path,
            "-vn",                 # No video
            "-map_metadata", "-1", # Remove metadata
            "-ac", "1",            # Convert to mono
            "-c:a", "libopus",     # Use Opus codec
            "-b:a", "12k",         # 12kbps bitrate
            "-application", "voip", # Optimize for voice
            output_path
        ], check=True, capture_output=True)
        
        # Check if compression was successful
        if os.path.exists(output_path) and os.path.getsize(output_path) > 0:
            return output_path
        else:
            return input_path
            
    except subprocess.SubprocessError:
        st.warning("Audio compression failed. Processing original file.")
        return input_path

def check_file_size(file_path):
    """Check if file size is under the API limit (25MB)"""
    max_size = 25 * 1024 * 1024  # 25MB in bytes
    return os.path.getsize(file_path) <= max_size

def transcribe_audio(file_path, language):
    """
    Transcribe audio using OpenAI Whisper API
    
    Args:
        file_path: Path to the audio file
        language: Language code (en or sv)
    
    Returns:
        Transcript text
    """
    client = openai.Client(api_key=get_openai_key())
    
    with open(file_path, "rb") as audio_file:
        transcript = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file,
            language=language
        )
    
    return transcript.text

def process_uploaded_file(uploaded_file, language_option, enable_diarization):
    """Process the uploaded audio or video file"""
    
    # Map language selection to language code
    language_map = {
        "English": "en",
        "Swedish": "sv"
    }
    language_code = language_map.get(language_option, "en")
    
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=Path(uploaded_file.name).suffix) as tmp_file:
        # Write the uploaded file to the temporary file
        tmp_file.write(uploaded_file.getbuffer())
        original_file_path = tmp_file.name
    
    try:
        file_size = os.path.getsize(original_file_path)
        file_size_mb = file_size / (1024 * 1024)
        
        # Check file size and compress if needed
        compressed_file_path = original_file_path
        if file_size_mb > 20:  # Compress if larger than 20MB as a safety margin
            st.info(f"File size: {file_size_mb:.2f}MB. Compressing for optimal processing...")
            compressed_file_path = compress_audio(original_file_path)
            compressed_size_mb = os.path.getsize(compressed_file_path) / (1024 * 1024)
            st.info(f"Compressed file size: {compressed_size_mb:.2f}MB")
            
        # Check if the file is still too large after compression
        if not check_file_size(compressed_file_path):
            st.error(f"File is too large for the API (max 25MB). Please try with a smaller file or contact support.")
            # Clean up
            if os.path.exists(original_file_path):
                os.remove(original_file_path)
            if compressed_file_path != original_file_path and os.path.exists(compressed_file_path):
                os.remove(compressed_file_path)
            return False
            
        # Transcribe the audio
        with st.spinner("Transcribing your file... This may take a moment."):
            transcript = transcribe_audio(compressed_file_path, language_code)
        
        # Store transcript in session state
        st.session_state.transcript = transcript
        st.session_state.file_name = uploaded_file.name
        
        # Perform speaker diarization if enabled
        if enable_diarization:
            with st.spinner("Performing speaker diarization... This may take a moment."):
                # Convert to WAV format for diarization if needed
                audio_for_diarization = compressed_file_path
                file_ext = Path(compressed_file_path).suffix.lower()
                
                # Convert video files or non-WAV audio to WAV for diarization
                if file_ext not in ['.wav']:
                    try:
                        temp_wav = os.path.join(tempfile.gettempdir(), f"diarize_{Path(compressed_file_path).stem}.wav")
                        subprocess.run([
                            "ffmpeg", 
                            "-i", compressed_file_path,
                            "-vn",       # No video
                            "-ac", "1",  # Mono
                            "-ar", "16000", # 16kHz sample rate
                            temp_wav
                        ], check=True, capture_output=True)
                        audio_for_diarization = temp_wav
                    except Exception as e:
                        st.error(f"Error converting to WAV format: {str(e)}")
                
                diarizer = Diarization()
                diarization_result = diarizer.process_audio(audio_for_diarization)
                
                # Clean up temporary WAV file if created
                if audio_for_diarization != compressed_file_path and os.path.exists(audio_for_diarization):
                    os.remove(audio_for_diarization)
                
                if "error" in diarization_result:
                    st.error(f"Diarization error: {diarization_result['error']}")
                    st.session_state.diarization_result = None
                    st.session_state.aligned_transcript = None
                else:
                    # Store diarization result in session state
                    st.session_state.diarization_result = diarization_result
                    
                    # Align transcript with diarization
                    aligned_segments = diarizer.align_transcript_with_diarization(
                        transcript, diarization_result
                    )
                    
                    # Format transcript with speaker labels
                    aligned_transcript = diarizer.format_transcript_with_speakers(aligned_segments)
                    st.session_state.aligned_transcript = aligned_transcript
        else:
            st.session_state.diarization_result = None
            st.session_state.aligned_transcript = None
        
        st.session_state.transcription_completed = True
        
        # Clean up the temporary files
        if os.path.exists(original_file_path):
            os.remove(original_file_path)
        if compressed_file_path != original_file_path and os.path.exists(compressed_file_path):
            os.remove(compressed_file_path)
        
        return True
    
    except Exception as e:
        st.error(f"An error occurred during processing: {str(e)}")
        # Clean up the temporary files
        if os.path.exists(original_file_path):
            os.remove(original_file_path)
        if compressed_file_path != original_file_path and os.path.exists(compressed_file_path):
            os.remove(compressed_file_path)
        return False

def export_to_docx(text, filename):
    """Export transcript to a docx file and create a download link"""
    
    # Create a new Document
    doc = Document()
    doc.add_heading('Transcript', 0)
    
    # Split the text into segments by speaker
    lines = text.split('\n')
    
    # Create a table to properly align text
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    
    # Configure table to hide borders
    from docx.oxml.shared import OxmlElement, qn
    
    # Function to set cell border properties
    def set_cell_border(cell, **kwargs):
        """Set cell border style"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # Get or create cell borders element
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ['start', 'top', 'end', 'bottom', 'insideH', 'insideV']:
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)
                element = OxmlElement(tag)
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))
                tcBorders.append(element)
                
        tcPr.append(tcBorders)
    
    # Define border style (zero width - effectively hidden)
    border_kwargs = {
        'start': {'sz': 0, 'val': 'nil'},
        'top': {'sz': 0, 'val': 'nil'},
        'end': {'sz': 0, 'val': 'nil'},
        'bottom': {'sz': 0, 'val': 'nil'},
        'insideH': {'sz': 0, 'val': 'nil'},
        'insideV': {'sz': 0, 'val': 'nil'}
    }
    
    # Set column widths
    from docx.shared import Inches
    table.columns[0].width = Inches(1.5)  # Width for speaker ID column
    table.columns[1].width = Inches(5.0)  # Width for text column
    
    # Add document styles
    from docx.shared import Pt
    normal_font = doc.styles['Normal']
    normal_font.font.name = 'Courier New'
    normal_font.font.size = Pt(10)
    
    # Variables to track current speaker and combined text
    current_speaker = None
    current_text = []
    
    # Process all lines
    for line in lines:
        line = line.strip()
        if not line or line == "Transcription Result:":
            continue  # Skip empty lines and the header
            
        # Check if this is a speaker line or continuation
        if line.startswith("SPEAKER_"):
            # If we have accumulated text for the previous speaker, add it now
            if current_speaker and current_text:
                row = table.add_row()
                speaker_cell = row.cells[0]
                text_cell = row.cells[1]
                
                # Add speaker info to first column
                p = speaker_cell.paragraphs[0]
                p.style = normal_font
                run = p.add_run(current_speaker)
                run.bold = True
                
                # Add text to second column
                p = text_cell.paragraphs[0]
                p.style = normal_font
                p.add_run('\n'.join(current_text))
                
                # Hide borders
                set_cell_border(speaker_cell, **border_kwargs)
                set_cell_border(text_cell, **border_kwargs)
                
                # Reset for new speaker
                current_text = []
            
            # Extract speaker and text
            parts = line.split(":", 1)
            if len(parts) == 2:
                current_speaker = parts[0] + ":"
                text = parts[1].strip()
                if text:
                    current_text.append(text)
        else:
            # Continuation line
            if current_text is not None:
                current_text.append(line)
    
    # Add the last speaker's text
    if current_speaker and current_text:
        row = table.add_row()
        speaker_cell = row.cells[0]
        text_cell = row.cells[1]
        
        p = speaker_cell.paragraphs[0]
        p.style = normal_font
        run = p.add_run(current_speaker)
        run.bold = True
        
        p = text_cell.paragraphs[0]
        p.style = normal_font
        p.add_run('\n'.join(current_text))
        
        # Hide borders
        set_cell_border(speaker_cell, **border_kwargs)
        set_cell_border(text_cell, **border_kwargs)
    
    # Save the document to a temporary file
    docx_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    docx_path = docx_file.name
    docx_file.close()  # Close the file handle immediately
    
    # Now save the document to the closed file
    doc.save(docx_path)
    
    # Read the saved file
    with open(docx_path, "rb") as file:
        docx_bytes = file.read()
    
    # Clean up the temporary file - use try/except to handle any permission errors
    try:
        os.unlink(docx_path)
    except PermissionError:
        # If we can't delete now, try to mark for deletion on exit
        try:
            import stat
            os.chmod(docx_path, stat.S_IWRITE)
            os.remove(docx_path)
        except:
            # If all else fails, leave it for cleanup later
            pass
    
    # Create download link
    b64 = base64.b64encode(docx_bytes).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}.docx">Download DOCX file</a>'
    
    return href

def main():
    """Main function for the Streamlit app"""
    
    # Password Protection
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False
        
    if not st.session_state.authenticated:
        st.title("Vocatio Lite - Login")
        password = st.text_input("Enter Password", type="password")
        
        # Get correct password from environment variables
        correct_password = os.environ.get("APP_PASSWORD")
        
        if not correct_password:
            st.error("No password has been set up. Add APP_PASSWORD to your environment variables or secrets.")
            return
            
        if st.button("Login"):
            if password == correct_password:
                st.session_state.authenticated = True
                st.experimental_rerun()
            else:
                st.error("Incorrect password")
        return
    
    # Main App (only shown if authenticated)
    # Title and description
    st.title("Vocatio Lite")
    st.write("Upload an audio or video file and get a transcript with speaker diarization.")
    
    # Check OpenAI API key
    api_key = get_openai_key()
    if not api_key:
        st.error("❌ OpenAI API key not found. Please set the OPENAI_API_KEY environment variable.")
        return
    
    # Check for ffmpeg (optional)
    try:
        subprocess.run(["ffmpeg", "-version"], check=True, capture_output=True)
    except (subprocess.SubprocessError, FileNotFoundError):
        st.warning("⚠️ ffmpeg not detected. Large files may not be processed correctly. Consider installing ffmpeg for better performance.")
    
    # File upload section
    st.subheader("Upload File")
    uploaded_file = st.file_uploader(
        "Choose an audio or video file", 
        type=["wav", "mp3", "mp4", "m4a", "ogg", "flac", "webm"]
    )
    
    # Language selection
    language_option = st.selectbox(
        "Select language",
        options=["English", "Swedish"],
        index=1
    )
    
    # Diarization option
    enable_diarization = st.checkbox("Enable speaker diarization", value=True, 
                                    help="Identify and separate different speakers in the audio")
    
    # Use token from environment variable only (no UI input)
    hf_token = get_hf_token()
    
    # Process the uploaded file
    if uploaded_file is not None and not st.session_state.transcription_completed:
        if st.button("Transcribe"):
            process_uploaded_file(uploaded_file, language_option, enable_diarization)
    
    # Display transcription results
    if st.session_state.transcription_completed:
        st.subheader("Transcription Result")
        
        # Display either the aligned transcript with speaker labels or the plain transcript
        display_text = st.session_state.aligned_transcript if st.session_state.aligned_transcript else st.session_state.transcript
        
        # Display as code for proper monospace formatting
        st.code(display_text, language=None)
        
        # Add a small note about the copy button
        st.caption("You can copy the transcript using the copy button in the top-right corner of the transcript box.")
        
        # Export as DOCX
        filename = os.path.splitext(st.session_state.file_name)[0] if st.session_state.file_name else "transcript"
        docx_download_link = export_to_docx(display_text, filename)
        st.markdown(docx_download_link, unsafe_allow_html=True)
        
        # Display speaker statistics if diarization was performed
        if st.session_state.diarization_result:
            st.subheader("Speaker Statistics")
            speaker_times = st.session_state.diarization_result["speaker_times"]
            
            # Calculate total duration
            total_duration = sum(speaker_times.values())
            
            # Create a table with speaker stats
            stats_data = []
            for speaker, duration in speaker_times.items():
                percentage = (duration / total_duration) * 100
                speaker_id = f"SPEAKER_{speaker}"
                stats_data.append({
                    "Speaker": speaker_id,
                    "Talk Time": f"{duration:.2f} seconds",
                    "Percentage": f"{percentage:.1f}%"
                })
            
            # Display as a table
            st.table(stats_data)
        
        # Reset button
        if st.button("Reset"):
            st.session_state.transcript = None
            st.session_state.transcription_completed = False
            st.session_state.file_name = None
            st.session_state.diarization_result = None
            st.session_state.aligned_transcript = None
            st.experimental_rerun()

if __name__ == "__main__":
    main() 